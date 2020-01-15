import os
from datetime import datetime
from docx import Document
from docx.shared import Pt

files = [f for f in os.listdir('.') if os.path.isfile(f)]

print('Select a template: ')
for i, doc in enumerate(files):
    print(i, doc)

input_index = int(input('Enter index: '))

input_file = files[input_index]


date = datetime.today().strftime('%Y/%m/%d')
recruiter_title = input("Enter recruiter title: ")
hook = input("Enter hook: ")

filein = open('../input.txt', 'r', encoding='utf-8')
text = filein.read()
lines = text.split('\n')

sal = ''
add2 = ''

for line in lines:
    if 'Organization:' in line:
        org = line[line.index('\t') + 1:]
    elif 'Salutation' in line:
        sal = line[line.index('\t') + 1:]
    elif 'Job Contact First Name' in line:
        first = line[line.index('\t') + 1:]
    elif 'Job Contact Last Name' in line:
        last = line[line.index('\t') + 1:]
    elif 'Address Line One' in line:
        add1 = line[line.index('\t') + 1:]
    elif 'Address Line Two' in line:
        add2 = line[line.index('\t') + 1:]
    elif 'City' in line:
        city = line[line.index('\t') + 1:]
    elif 'Province' in line:
        prov = line[line.index('\t') + 1:]
    elif 'Postal Code' in line:
        postal = line[line.index('\t') + 1:]
    elif 'Job ID:' in line:
        line = line.split(' ')
        position = ' '.join(line[3:])
        position_id = ' '.join(line[1:3])

name = f'{first} {last}'

if sal != '':
    to_name = f'{sal} {last}'
else:
    to_name = name

document = Document(input_file)

for i, para in enumerate(document.paragraphs):
    edited = False

    if "DATE" in para.text:
        document.paragraphs[i].text = para.text.replace("DATE", date)
        edited = True
    if "FULL_NAME" in para.text:
        document.paragraphs[i].text = para.text.replace("FULL_NAME", name)
        edited = True
    if "RECRUITER_TITLE" in para.text:
        document.paragraphs[i].text = para.text.replace("RECRUITER_TITLE", recruiter_title)
        edited = True
    if "RECRUITER_TITLE" in para.text:
        document.paragraphs[i].text = para.text.replace("RECRUITER_TITLE", sal)
        edited = True
    if "COMPANY" in para.text:
        document.paragraphs[i].text = para.text.replace("COMPANY", org)
        edited = True
    if "STREET_ADDRESS" in para.text:
        document.paragraphs[i].text = para.text.replace("STREET_ADDRESS", add1 + ' ' +add2)
        edited = True
    if "CITY" in para.text:
        document.paragraphs[i].text = para.text.replace("CITY", city)
        edited = True
    if "PROVINCE" in para.text:
        document.paragraphs[i].text = para.text.replace("PROVINCE", prov)
        edited = True
    if "POSTAL_CODE" in para.text:
        document.paragraphs[i].text = para.text.replace("POSTAL_CODE", postal)
        edited = True
    if "POSITION_NAME" in para.text:
        document.paragraphs[i].text = para.text.replace("POSITION_NAME", position)
        edited = True
    if "ID:" in para.text:
        document.paragraphs[i].text = para.text.replace("ID:", position_id)
        edited = True
    if "TO_NAME" in para.text:
        document.paragraphs[i].text = para.text.replace("TO_NAME", to_name)
        edited = True
    if "POSITION" in para.text:
        document.paragraphs[i].text = para.text.replace("POSITION", position)
        edited = True
    if "HOOK" in para.text:
        document.paragraphs[i].text = para.text.replace("HOOK", hook)
        edited = True
    
    if edited:
        document.paragraphs[i].style.font.name = 'Cambria'
        document.paragraphs[i].style.font.size = Pt(13)

document.save(f'../{org.lower()}-{input_file}'.replace(' ', '_'))
